import unittest
from io import BytesIO

from docx import Document

from utils.text_splitter import extract_text_from_docx_bytes, split_text


class TextSplitterTest(unittest.TestCase):
    def test_long_unpunctuated_text_is_capped_to_chunk_size(self):
        text = "Near limit phrase VIOLET HARBOR. " + ("data pattern " * 3900)

        chunks = split_text(text, chunk_size=1200, chunk_overlap=200)

        self.assertGreater(len(chunks), 1)
        self.assertLessEqual(max(len(chunk) for chunk in chunks), 1200)
        self.assertIn("VIOLET HARBOR", chunks[0])

    def test_extract_docx_text_includes_paragraphs_and_table_cells(self):
        document = Document()
        document.add_paragraph("Project codename: BLUE LANTERN")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Owner"
        table.cell(0, 1).text = "Aether Ops"

        buffer = BytesIO()
        document.save(buffer)

        text = extract_text_from_docx_bytes(buffer.getvalue())

        self.assertIn("BLUE LANTERN", text)
        self.assertIn("Owner", text)
        self.assertIn("Aether Ops", text)


if __name__ == "__main__":
    unittest.main()
