import re
from io import BytesIO


def split_text(
    text: str,
    chunk_size: int = 1200,
    chunk_overlap: int = 200,
) -> list[str]:
    """
    Split text into overlapping chunks at paragraph boundaries.
    Larger chunks preserve more context per retrieval hit.
    """
    # Normalise excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    # Split on paragraph boundaries first (double newline)
    paragraphs = text.split("\n\n")
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Paragraph fits in the current chunk
        if len(current_chunk) + len(para) + 2 <= chunk_size:
            current_chunk = (current_chunk + "\n\n" + para).strip()
        else:
            # Flush current chunk
            if current_chunk:
                chunks.append(current_chunk)

            # Paragraph itself is too long — split at sentence boundaries
            if len(para) > chunk_size:
                sub_chunks = _split_long_section(para, chunk_size, chunk_overlap)
                # Keep the last sub-chunk as the start of the next chunk
                if sub_chunks:
                    chunks.extend(sub_chunks[:-1])
                    current_chunk = sub_chunks[-1]
                else:
                    current_chunk = ""
            else:
                # Start new chunk with overlap from the previous one
                if chunk_overlap > 0 and current_chunk:
                    overlap_text = _get_overlap(current_chunk, chunk_overlap)
                    current_chunk = (overlap_text + "\n\n" + para).strip()
                else:
                    current_chunk = para

    if current_chunk:
        chunks.append(current_chunk)

    return [c for c in chunks if c.strip()]


def _get_overlap(text: str, overlap_size: int) -> str:
    """Return the last `overlap_size` characters of text, aligned to a word boundary."""
    if len(text) <= overlap_size:
        return text
    snippet = text[-overlap_size:]
    # Try to align to a word start
    space_pos = snippet.find(" ")
    if space_pos > 0:
        return snippet[space_pos:].strip()
    return snippet.strip()


def _split_long_section(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """Split a section that is too long, splitting at sentence boundaries."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks: list[str] = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) + 1 <= chunk_size:
            current = (current + " " + sentence).strip()
        else:
            if current:
                chunks.append(current)
                current = ""
            if len(sentence) > chunk_size:
                sub_chunks = _split_by_length(sentence, chunk_size, chunk_overlap)
                if sub_chunks:
                    chunks.extend(sub_chunks[:-1])
                    current = sub_chunks[-1]
                continue
            if chunk_overlap > 0 and current:
                overlap = _get_overlap(current, chunk_overlap)
                current = (overlap + " " + sentence).strip()
            else:
                current = sentence

    if current:
        chunks.append(current)

    return chunks


def _split_by_length(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """Split text that has no usable sentence boundaries."""
    words = text.split()
    if not words:
        return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    chunks: list[str] = []
    current = ""

    for word in words:
        if len(word) > chunk_size:
            if current:
                chunks.append(current)
                current = ""
            chunks.extend(word[i:i + chunk_size] for i in range(0, len(word), chunk_size))
            continue

        candidate = (current + " " + word).strip()
        if len(candidate) <= chunk_size:
            current = candidate
            continue

        if current:
            chunks.append(current)
            overlap = _get_overlap(current, chunk_overlap) if chunk_overlap > 0 else ""
            current = (overlap + " " + word).strip() if overlap else word
        else:
            current = word

    if current:
        chunks.append(current)

    return chunks


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).
    Handles columns, tables, and complex layouts far better than pypdf.
    """
    try:
        import fitz  # pymupdf
        import io

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text: list[str] = []

        for page in doc:
            # "text" mode preserves reading order across columns
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                pages_text.append(page_text.strip())

        doc.close()
        return "\n\n".join(pages_text)

    except ImportError:
        # Fallback to pypdf if pymupdf is somehow not installed
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(p.strip() for p in pages if p.strip())
        except Exception as e:
            raise ValueError(f"PDF extraction failed (no pymupdf): {e}")
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract visible paragraph and table text from a DOCX file."""
    try:
        from docx import Document

        document = Document(BytesIO(file_bytes))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts)

    except ImportError as e:
        raise ValueError(f"DOCX extraction failed (python-docx not installed): {e}")
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}")


def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
    """Decode plain-text uploads using UTF-8, then latin-1 as a permissive fallback."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            raise ValueError(f"Could not decode text file: {e}")
